import requests
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
import time
import os
import json
from selenium.common.exceptions import TimeoutException
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.shared import RGBColor
from datetime import date

urls_sri = [
    'https://srienlinea.sri.gob.ec/sri-en-linea/SriRucWeb/ConsultaRuc/Consultas/consultaRuc',
    'https://srienlinea.sri.gob.ec/sri-en-linea/SriDeclaracionesWeb/EstadoTributario/Consultas/consultaEstadoTributario',
    'https://srienlinea.sri.gob.ec/sri-en-linea/SriPagosWeb/ConsultaDeudasFirmesImpugnadas/Consultas/consultaDeudasFirmesImpugnadas'
]
url_aduana = 'https://www.aduana.gob.ec/servicio-al-ciudadano/consulta-de-certificado-cumplimiento/'
url_fiscalia = 'https://www.fiscalia.gob.ec/consulta-de-noticias-del-delito/'
url_consejo_judicatura = 'https://procesosjudiciales.funcionjudicial.gob.ec/busqueda-filtros'
docs_data = [
    "Consulta de RUC/ Empresas fantasmas del SRI",
    "Estado Tributario",
    "Deudas firmes e impugnadas"
]

output_dir = 'screenshots'
os.makedirs(output_dir, exist_ok=True)

def add_header_image(document, image_path):
    section = document.sections[0]
    section.page_width = Pt(595.276)
    section.page_height = Pt(841.890)

    section.header_distance = Inches(0)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    header = section.header
    header_para = header.paragraphs[0]
    header_para.paragraph_format.left_indent = Inches(-1)
    run = header_para.add_run()
    picture = run.add_picture(image_path, width=Inches(8.27))

    pic = picture._inline.graphic.graphicData.pic
    pic_pr = pic.spPr

    a = OxmlElement('a:effectLst')
    pic_pr.append(a)

    pic_pr.append(OxmlElement('a:relativeHeight'))
    pic_pr.append(OxmlElement('a:behindDoc'))

    return document

def format_main_heading(document, text):
    heading = document.add_heading(text, level=1)
    heading_paragraph = heading.paragraph_format
    heading_paragraph.alignment = 1
    heading_run = heading.runs[0]
    heading_run.font.name = 'Arial'
    heading_run.font.size = Pt(11)
    heading_run.font.bold = True
    heading_run.font.size = Pt(11)
    heading_run.font.color.rgb = RGBColor(0, 0, 0)
    return heading

def format_subheading(document, text):
    heading = document.add_heading(text, level=2)
    heading_paragraph = heading.paragraph_format
    heading_paragraph.alignment = 3
    heading_run = heading.runs[0]
    heading_run.font.name = 'Arial'
    heading_run.font.size = Pt(11)
    heading_run.font.bold = True
    heading_run.font.underline = True
    heading_run.font.color.rgb = RGBColor(0, 0, 0)
    return heading

def format_subheading2(document, text):
    heading = document.add_heading(text, level=2)
    heading_paragraph = heading.paragraph_format
    heading_paragraph.alignment = 3
    heading_run = heading.runs[0]
    heading_run.font.name = 'Arial'
    heading_run.font.size = Pt(11)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(0, 0, 0)
    return heading

def create_source_table(document, second_column_text):
    table = document.add_table(1, 2)
    table.cell(0, 0).text = 'Fuente:'
    table.cell(0, 1).text = second_column_text

def create_source_table1(document, second_column_text):
    table = document.add_table(2, 2)
    table.columns[0].width = Inches(0.6)
    table.cell(0, 0).text = 'Busqueda:'
    table.cell(1, 0).text = 'Fecha de Busqueda:'
    table.cell(0, 1).text = second_column_text
    table.cell(1, 1).text = str(date.today())

def create_source_table2(document):
    table = document.add_table(5, 2)
    table.cell(0, 0).text = 'Código de procedimiento'
    table.cell(1, 0).text = 'Tipo de procedimiento'
    table.cell(2, 0).text = 'Objeto contractual'
    table.cell(3, 0).text = 'Presupuesto'
    table.cell(4, 0).text = 'Estado a la fecha de elaboración'

def create_source_table3(document, ruc):
    table = document.add_table(7, 2)
    table.cell(0, 0).text = 'Nombre'
    table.cell(1, 0).text = 'RUC'
    table.cell(2, 0).text = 'Participación en el procedimiento'
    table.cell(3, 0).text = 'Domicilio Fiscal'
    table.cell(4, 0).text = 'Fecha inicio de actividades'
    table.cell(5, 0).text = 'Actividad económica principal'
    table.cell(6, 0).text = 'Recurrencias'
    data = get_ruc_data(ruc)
    table.cell(0, 1).text = data[0]['razonSocial']
    table.cell(1, 1).text = str(ruc)
    table.cell(4, 1).text = data[0]['informacionFechasContribuyente']['fechaInicioActividades']
    table.cell(5, 1).text = data[0]['actividadEconomicaPrincipal']

def create_source_table4(document):
    table = document.add_table(9, 5)
    table.cell(0, 0).text = 'N°'
    table.cell(0, 1).text = 'Riesgos potenciales'
    table.cell(0, 2).text = 'Hallazgos'
    table.cell(0, 3).text = 'Página(s)'
    table.cell(1, 0).text = '1'
    table.cell(2, 0).text = '2'
    table.cell(3, 0).text = '3'
    table.cell(4, 0).text = '4'
    table.cell(5, 0).text = '5'
    table.cell(6, 0).text = '6'
    table.cell(7, 0).text = '7'
    table.cell(8, 0).text = '8'
    table.cell(1, 1).text = 'Soborno, corrupción y fraude'
    table.cell(2, 1).text = 'Relaciones gubernamentales'
    table.cell(3, 1).text = 'Actividades criminales/ilegales'
    table.cell(4, 1).text = 'Incumplimiento financiero'
    table.cell(5, 1).text = 'Asuntos regulatorios'
    table.cell(6, 1).text = 'Litigios'
    table.cell(7, 1).text = 'Medios adversos'
    table.cell(8, 1).text = 'Otros asuntos'

def add_format(document, ruc):

    add_header_image(document, 'assets/img.png')
    format_main_heading(document, 'Formato para documentar las Debidas Diligencias Proveedores')
    format_subheading(document, 'Resumen ejecutivo')
    paragraph = document.add_paragraph(
        'Debida Diligencia que se realiza en cumplimiento a la ejecución de controles del punto 8.2 del Manual Operativo del Sistema de Gestión Antisoborno, alineado a la Norma ISO 37001:2016 y controles de la Matriz de Riesgos.')
    paragraph_style = paragraph.paragraph_format
    paragraph_style.alignment = 3
    format_subheading2(document, 'Información del procedimiento de contratación asociado:')
    create_source_table2(document)
    format_subheading2(document, 'Información del procedimiento de contratación asociado:')
    create_source_table3(document, ruc)
    create_source_table4(document)
    create_source_table(document, 'Servicio de Rentas Internas')

# options.add_argument('--headless')
# options.add_argument('--disable-gpu')
# options.add_argument('--no-sandbox')
options = Options()
options.add_argument('--window-size=1920,1080')

def validate_ruc(ruc):
    validation_url = f"https://srienlinea.sri.gob.ec/sri-catastro-sujeto-servicio-internet/rest/ConsolidadoContribuyente/existePorNumeroRuc?numeroRuc={ruc}"
    response = requests.get(validation_url)
    return response.content.decode('utf8')

def get_ruc_data(ruc):
    data_url = f"https://srienlinea.sri.gob.ec/sri-catastro-sujeto-servicio-internet/rest/ConsolidadoContribuyente/obtenerPorNumerosRuc?&ruc={ruc}"
    response = requests.get(data_url)
    return response.json()

def scrape_from_sri(document, ruc):
    for i, url in enumerate(urls_sri):
        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
        try:
            driver.get(url)
            try:
                ruc_input = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.ID, "busquedaRucId"))
                )
                print("Input field found")
                ruc_input.clear()
                ruc_input.send_keys(ruc)
            except Exception as e:
                print(f"Error finding input field: {str(e)}")
                raise
            try:
                button = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, "//button[contains(@class, 'ui-button-text-only') and .//span[contains(text(), 'Consultar')]]"))
                )
                print("Button found")
                button.click()
                print("Button clicked")
            except Exception as e:
                print(f"Error with button interaction: {str(e)}")
                raise
            time.sleep(20)
            screenshot_path = os.path.join(output_dir, f'capture_{i + 1}.png')
            driver.save_screenshot(screenshot_path)

        except Exception as e:
            screenshot_path = os.path.join(output_dir, f'capture_{i + 1}.png')
            driver.save_screenshot(screenshot_path)

        create_source_table1(document, docs_data[i])
        document.add_picture(screenshot_path, width=Inches(6))
        document.add_paragraph('')

        driver.quit()

def scrape_from_aduana(document, ruc):
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
    try:
        driver.get(url_aduana)
        try:
            ruc_input = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.ID, "txtRuc"))
            )
            print("Input field found")
            ruc_input.clear()
            ruc_input.send_keys(ruc)
        except Exception as e:
            print(f"Error finding input field: {str(e)}")
            raise
        try:
            button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.ID, 'btnConsultar'))
            )
            print("Button found")
            button.click()
            print("Button clicked")
        except Exception as e:
            print(f"Error with button interaction: {str(e)}")
            raise
        time.sleep(15)
        screenshot_path = os.path.join(output_dir, f'capture_4.png')
        driver.save_screenshot(screenshot_path)

    except Exception as e:
        screenshot_path = os.path.join(output_dir, f'capture_4.png')
        driver.save_screenshot(screenshot_path)

    create_source_table1(document, 'Liquidaciones vencidas')
    document.add_picture(screenshot_path, width=Inches(6))
    document.add_paragraph('')

    driver.quit()

def scrape_from_fiscalia(document, ruc):
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
    try:
        driver.get(url_fiscalia)
        time.sleep(10)
        try:
            iframe = WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.ID, "blockrandom"))
            )
            driver.switch_to.frame(iframe)
            print("Switched to iframe successfully")
        except Exception as e:
            print(f"Error switching to iframe: {str(e)}")
            raise
        try:
            ruc_input = WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.ID, "pwd"))
            )
            print("Input field found")
            ruc_input.clear()
            ruc_input.send_keys(ruc)
        except Exception as e:
            print(f"Error finding input field: {str(e)}")
            raise
        try:
            button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.ID, 'btn_buscar_denuncia'))
            )
            print("Button found")
            button.click()
            print("Button clicked")
        except Exception as e:
            print(f"Error with button interaction: {str(e)}")
            raise

        driver.switch_to.default_content()
        
        time.sleep(10)
        screenshot_path = os.path.join(output_dir, f'capture_5.png')
        driver.save_screenshot(screenshot_path)

    except Exception as e:
        print(f"General error in scrape_from_fiscalia: {str(e)}")
        screenshot_path = os.path.join(output_dir, f'capture_5.png')
        driver.save_screenshot(screenshot_path)

    create_source_table1(document, 'Procesos Fiscales')
    document.add_picture(screenshot_path, width=Inches(6))
    document.add_paragraph('')

    driver.quit()

def scrape_from_consejo_judicatura(document, ruc):
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
    try:
        driver.get(url_consejo_judicatura)
        try:
            ruc_input = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.ID, "mat-input-3"))
            )
            print("Input field found")
            ruc_input.clear()
            ruc_input.send_keys(ruc)
        except Exception as e:
            print(f"Error finding input field: {str(e)}")
            raise
        time.sleep(10)
        try:
            button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, "button.boton-buscar"))
            )
            print("Button found")
            button.click()
            print("Button clicked")
        except Exception as e:
            print(f"Error with button interaction: {str(e)}")
            raise
        time.sleep(15)
        try:
            button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, "button.boton-buscar"))
            )
            print("Button found")
            button.click()
            print("Button clicked")
        except Exception as e:
            print(f"Error with button interaction: {str(e)}")
            raise
        time.sleep(5)
        screenshot_path = os.path.join(output_dir, f'capture_6.png')
        driver.save_screenshot(screenshot_path)

    except Exception as e:
        screenshot_path = os.path.join(output_dir, f'capture_6.png')
        driver.save_screenshot(screenshot_path)

    create_source_table1(document, 'Liquidaciones vencidas')
    document.add_picture(screenshot_path, width=Inches(6))
    document.add_paragraph('')

    driver.quit()

if __name__ == '__main__':
    ruc = "0190340325001"
    if validate_ruc(ruc) == 'true':
        document = Document()
        add_format(document, ruc)
        scrape_from_sri(document, ruc)
        create_source_table(document, 'Aduana del Ecuador')
        scrape_from_aduana(document, ruc)
        create_source_table(document, 'Fiscalía')
        scrape_from_fiscalia(document, ruc)
        create_source_table(document, 'Consejo de la Judicatura')
        scrape_from_consejo_judicatura(document, ruc)
        document.save(os.path.join(output_dir, 'evidence_report.docx'))
        print("Evidence report saved as 'evidence_report.docx'")
    else:
        print("RUC no valido")
