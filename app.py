import requests
from selenium import webdriver
from selenium.webdriver.firefox.options import Options
from selenium.webdriver.firefox.service import Service
from webdriver_manager.firefox import GeckoDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
import time
import os
import json
from selenium.common.exceptions import TimeoutException
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.shared import RGBColor
from datetime import date
from flask import Flask, render_template, request, flash, redirect, url_for
from flask import send_file

urls_sri = [
    'https://srienlinea.sri.gob.ec/sri-en-linea/SriRucWeb/ConsultaRuc/Consultas/consultaRuc',
    'https://srienlinea.sri.gob.ec/sri-en-linea/SriDeclaracionesWeb/EstadoTributario/Consultas/consultaEstadoTributario',
    'https://srienlinea.sri.gob.ec/sri-en-linea/SriPagosWeb/ConsultaDeudasFirmesImpugnadas/Consultas/consultaDeudasFirmesImpugnadas'
]
url_aduana = 'https://www.aduana.gob.ec/servicio-al-ciudadano/consulta-de-certificado-cumplimiento/'
url_fiscalia = 'https://www.fiscalia.gob.ec/consulta-de-noticias-del-delito/'
url_consejo_judicatura = 'https://procesosjudiciales.funcionjudicial.gob.ec/busqueda-filtros'
url_soce_incumplidos = 'https://www.compraspublicas.gob.ec/ProcesoContratacion/compras/EP/EmpReporteIncumplidos.cpe'
url_contraloria = 'https://www.contraloria.gob.ec/Consultas/InformesAprobados'
url_senescyt = 'https://www.senescyt.gob.ec/consulta-titulos-web/faces/vista/consulta/consulta.xhtml'
docs_data = [
    "Consulta de RUC/ Empresas fantasmas del SRI",
    "Estado Tributario",
    "Deudas firmes e impugnadas"
]
output_dir = 'screenshots'
os.makedirs(output_dir, exist_ok=True)

def add_header_image(document, image_path):
    section = document.sections[0]
    section.page_width = Pt(595.276)
    section.page_height = Pt(841.890)

    section.header_distance = Inches(0)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    header = section.header
    header_para = header.paragraphs[0]
    header_para.paragraph_format.left_indent = Inches(-1)
    run = header_para.add_run()
    picture = run.add_picture(image_path, width=Inches(8.27))

    pic = picture._inline.graphic.graphicData.pic
    pic_pr = pic.spPr

    a = OxmlElement('a:effectLst')
    pic_pr.append(a)

    pic_pr.append(OxmlElement('a:relativeHeight'))
    pic_pr.append(OxmlElement('a:behindDoc'))

    return document

def format_main_heading(document, text):
    heading = document.add_heading(text, level=1)
    heading_paragraph = heading.paragraph_format
    heading_paragraph.alignment = 1
    heading_run = heading.runs[0]
    heading_run.font.name = 'Arial'
    heading_run.font.size = Pt(11)
    heading_run.font.bold = True
    heading_run.font.size = Pt(11)
    heading_run.font.color.rgb = RGBColor(0, 0, 0)
    return heading

def format_subheading(document, text):
    heading = document.add_heading(text, level=2)
    heading_paragraph = heading.paragraph_format
    heading_paragraph.alignment = 3
    heading_run = heading.runs[0]
    heading_run.font.name = 'Arial'
    heading_run.font.size = Pt(11)
    heading_run.font.bold = True
    heading_run.font.underline = True
    heading_run.font.color.rgb = RGBColor(0, 0, 0)
    return heading

def format_subheading2(document, text):
    heading = document.add_heading(text, level=2)
    heading_paragraph = heading.paragraph_format
    heading_paragraph.alignment = 3
    heading_run = heading.runs[0]
    heading_run.font.name = 'Arial'
    heading_run.font.size = Pt(11)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(0, 0, 0)
    return heading

def create_source_table(document, second_column_text):
    table = document.add_table(1, 2)
    table.cell(0, 0).text = 'Fuente:'
    table.cell(0, 1).text = second_column_text

def create_source_table1(document, second_column_text):
    table = document.add_table(2, 2)
    table.columns[0].width = Inches(0.6)
    table.cell(0, 0).text = 'Busqueda:'
    table.cell(1, 0).text = 'Fecha de Busqueda:'
    table.cell(0, 1).text = second_column_text
    table.cell(1, 1).text = str(date.today())

def create_source_table2(document):
    table = document.add_table(5, 2)
    table.cell(0, 0).text = 'Código de procedimiento'
    table.cell(1, 0).text = 'Tipo de procedimiento'
    table.cell(2, 0).text = 'Objeto contractual'
    table.cell(3, 0).text = 'Presupuesto'
    table.cell(4, 0).text = 'Estado a la fecha de elaboración'

def create_source_table3(document, ruc):
    table = document.add_table(7, 2)
    table.cell(0, 0).text = 'Nombre'
    table.cell(1, 0).text = 'RUC'
    table.cell(2, 0).text = 'Participación en el procedimiento'
    table.cell(3, 0).text = 'Domicilio Fiscal'
    table.cell(4, 0).text = 'Fecha inicio de actividades'
    table.cell(5, 0).text = 'Actividad económica principal'
    table.cell(6, 0).text = 'Recurrencias'
    data = get_ruc_data(ruc)
    table.cell(0, 1).text = data[0]['razonSocial']
    table.cell(1, 1).text = str(ruc)
    table.cell(4, 1).text = data[0]['informacionFechasContribuyente']['fechaInicioActividades']
    table.cell(5, 1).text = data[0]['actividadEconomicaPrincipal']

def create_source_table4(document):
    table = document.add_table(9, 5)
    table.cell(0, 0).text = 'N°'
    table.cell(0, 1).text = 'Riesgos potenciales'
    table.cell(0, 2).text = 'Hallazgos'
    table.cell(0, 3).text = 'Página(s)'
    table.cell(1, 0).text = '1'
    table.cell(2, 0).text = '2'
    table.cell(3, 0).text = '3'
    table.cell(4, 0).text = '4'
    table.cell(5, 0).text = '5'
    table.cell(6, 0).text = '6'
    table.cell(7, 0).text = '7'
    table.cell(8, 0).text = '8'
    table.cell(1, 1).text = 'Soborno, corrupción y fraude'
    table.cell(2, 1).text = 'Relaciones gubernamentales'
    table.cell(3, 1).text = 'Actividades criminales/ilegales'
    table.cell(4, 1).text = 'Incumplimiento financiero'
    table.cell(5, 1).text = 'Asuntos regulatorios'
    table.cell(6, 1).text = 'Litigios'
    table.cell(7, 1).text = 'Medios adversos'
    table.cell(8, 1).text = 'Otros asuntos'

def create_source_table5(document, ruc):
    data = get_ruc_data(ruc)
    table = document.add_table(4, 2)
    table.cell(0, 0).text = 'Nombre'
    table.cell(1, 0).text = 'Identificación'
    table.cell(2, 0).text = 'Cargo'
    table.cell(3, 0).text = 'Nacionalidad'
    table.cell(0, 1).text = data[0]['representantesLegales'][0]['nombre']
    identification = data[0]['representantesLegales'][0]['identificacion']
    table.cell(1, 1).text = identification
    return identification

def create_source_table6(document):
    table = document.add_table(9, 5)
    table.cell(0, 0).text = 'N°'
    table.cell(0, 1).text = 'Riesgos potenciales'
    table.cell(0, 2).text = 'Hallazgos'
    table.cell(0, 3).text = 'Página(s)'
    table.cell(1, 0).text = '1'
    table.cell(2, 0).text = '2'
    table.cell(3, 0).text = '3'
    table.cell(4, 0).text = '4'
    table.cell(5, 0).text = '5'
    table.cell(6, 0).text = '6'
    table.cell(7, 0).text = '7'
    table.cell(8, 0).text = '8'
    table.cell(1, 1).text = 'Soborno, corrupción y fraude'
    table.cell(2, 1).text = 'Relaciones gubernamentales'
    table.cell(3, 1).text = 'Actividades criminales/ilegales'
    table.cell(4, 1).text = 'Incumplimiento financiero'
    table.cell(5, 1).text = 'Asuntos regulatorios'
    table.cell(6, 1).text = 'Litigios'
    table.cell(7, 1).text = 'Medios adversos'
    table.cell(8, 1).text = 'Otros asuntos'

def add_format(document, ruc):

    add_header_image(document, 'assets/img.png')
    format_main_heading(document, 'Formato para documentar las Debidas Diligencias Proveedores')
    format_subheading(document, 'Resumen ejecutivo')
    paragraph = document.add_paragraph(
        'Debida Diligencia que se realiza en cumplimiento a la ejecución de controles del punto 8.2 del Manual Operativo del Sistema de Gestión Antisoborno, alineado a la Norma ISO 37001:2016 y controles de la Matriz de Riesgos.')
    paragraph_style = paragraph.paragraph_format
    paragraph_style.alignment = 3
    format_subheading2(document, 'Información del procedimiento de contratación asociado:')
    create_source_table2(document)
    format_subheading2(document, 'Información del procedimiento de contratación asociado:')
    create_source_table3(document, ruc)
    create_source_table4(document)

def add_format2(document, ruc):
    format_subheading(document, 'Representante Legal')
    format_subheading2(document, 'Información del sujeto a revisión:')
    identification = create_source_table5(document, ruc)
    create_source_table6(document)
    return identification

options = Options()
# options.add_argument('--headless')
options.add_argument('--width=1920')
options.add_argument('--height=1080')

def validate_ruc(ruc):
    validation_url = f"https://srienlinea.sri.gob.ec/sri-catastro-sujeto-servicio-internet/rest/ConsolidadoContribuyente/existePorNumeroRuc?numeroRuc={ruc}"
    response = requests.get(validation_url)
    return response.content.decode('utf8')

def get_ruc_data(ruc):
    data_url = f"https://srienlinea.sri.gob.ec/sri-catastro-sujeto-servicio-internet/rest/ConsolidadoContribuyente/obtenerPorNumerosRuc?&ruc={ruc}"
    response = requests.get(data_url)
    return response.json()

def scrape_from_sri(document, ruc):
    for i, url in enumerate(urls_sri):
        driver = webdriver.Firefox(service=Service(GeckoDriverManager().install()), options=options)
        try:
            driver.get(url)
            try:
                ruc_input = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.ID, "busquedaRucId"))
                )
                print("Input field found")
                ruc_input.clear()
                ruc_input.send_keys(ruc)
            except Exception as e:
                print(f"Error finding input field: {str(e)}")
                raise
                
            try:
                button = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, "//button[contains(@class, 'ui-button-text-only') and .//span[contains(text(), 'Consultar')]]"))
                )
                print("Button found")
                button.click()
                print("Button clicked")
                
                # Modified CAPTCHA handling
                try:
                    # First try to find the reCAPTCHA iframe
                    iframes = WebDriverWait(driver, 10).until(
                        EC.presence_of_all_elements_located((By.TAG_NAME, "iframe"))
                    )
                    
                    captcha_iframe = None
                    for iframe in iframes:
                        if 'reCAPTCHA' in iframe.get_attribute('title') or 'recaptcha' in iframe.get_attribute('src'):
                            captcha_iframe = iframe
                            break
                    
                    if captcha_iframe:
                        print("CAPTCHA iframe found")
                        driver.switch_to.frame(captcha_iframe)
                        
                        # Wait for either image grid or checkbox
                        try:
                            WebDriverWait(driver, 5).until(
                                EC.presence_of_element_located((By.CLASS_NAME, "rc-imageselect-target"))
                            )
                            print("Image CAPTCHA detected")
                        except:
                            try:
                                checkbox = WebDriverWait(driver, 5).until(
                                    EC.presence_of_element_located((By.CLASS_NAME, "recaptcha-checkbox-border"))
                                )
                                checkbox.click()
                                print("Checkbox CAPTCHA clicked")
                            except:
                                print("No standard CAPTCHA elements found")
                        
                        # Wait for manual solving
                        print("Waiting for manual CAPTCHA solution...")
                        WebDriverWait(driver, 60).until(
                            lambda x: 'rc-imageselect-error-select-more' not in driver.page_source
                        )
                        
                        # Switch back to main content
                        driver.switch_to.default_content()
                        print("CAPTCHA handling completed")
                    
                except Exception as e:
                    print(f"CAPTCHA handling error: {str(e)}")
                    driver.switch_to.default_content()
                
            except Exception as e:
                print(f"Error with button interaction: {str(e)}")
                raise
                
            # Wait for results to load after CAPTCHA
            time.sleep(20)
            screenshot_path = os.path.join(output_dir, 'capture.png')
            driver.save_screenshot(screenshot_path)

        except Exception as e:
            screenshot_path = os.path.join(output_dir, 'capture.png')
            driver.save_screenshot(screenshot_path)

        create_source_table1(document, docs_data[i])
        document.add_picture(screenshot_path, width=Inches(6))
        document.add_paragraph('')

        driver.quit()

def scrape_from_aduana(document, ruc):
    driver = webdriver.Chrome(service=Service(GeckoDriverManager().install()), options=options)
    try:
        driver.get(url_aduana)
        try:
            ruc_input = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.ID, "txtRuc"))
            )
            print("Input field found")
            ruc_input.clear()
            ruc_input.send_keys(ruc)
        except Exception as e:
            print(f"Error finding input field: {str(e)}")
            raise
        try:
            button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.ID, 'btnConsultar'))
            )
            print("Button found")
            button.click()
            print("Button clicked")
        except Exception as e:
            print(f"Error with button interaction: {str(e)}")
            raise
        time.sleep(15)
        screenshot_path = os.path.join(output_dir, 'capture.png')
        driver.save_screenshot(screenshot_path)

    except Exception as e:
        screenshot_path = os.path.join(output_dir, 'capture.png')
        driver.save_screenshot(screenshot_path)

    create_source_table1(document, 'Liquidaciones vencidas')
    document.add_picture(screenshot_path, width=Inches(6))
    document.add_paragraph('')

    driver.quit()

def scrape_from_fiscalia(document, ruc):
    driver = webdriver.Chrome(service=Service(GeckoDriverManager().install()), options=options)
    try:
        driver.get(url_fiscalia)
        time.sleep(10)
        try:
            iframe = WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.ID, "blockrandom"))
            )
            driver.switch_to.frame(iframe)
            print("Switched to iframe successfully")
        except Exception as e:
            print(f"Error switching to iframe: {str(e)}")
            raise
        try:
            ruc_input = WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.ID, "pwd"))
            )
            print("Input field found")
            ruc_input.clear()
            ruc_input.send_keys(ruc)
        except Exception as e:
            print(f"Error finding input field: {str(e)}")
            raise
        try:
            button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.ID, 'btn_buscar_denuncia'))
            )
            print("Button found")
            button.click()
            print("Button clicked")
        except Exception as e:
            print(f"Error with button interaction: {str(e)}")
            raise

        driver.switch_to.default_content()
        
        time.sleep(10)
        screenshot_path = os.path.join(output_dir, 'capture.png')
        driver.save_screenshot(screenshot_path)

    except Exception as e:
        print(f"General error in scrape_from_fiscalia: {str(e)}")
        screenshot_path = os.path.join(output_dir, 'capture.png')
        driver.save_screenshot(screenshot_path)

    create_source_table1(document, 'Procesos Fiscales')
    document.add_picture(screenshot_path, width=Inches(6))
    document.add_paragraph('')

    driver.quit()

def scrape_from_consejo_judicatura(document, ruc):
    driver = webdriver.Chrome(service=Service(GeckoDriverManager().install()), options=options)
    try:
        driver.get(url_consejo_judicatura)
        try:
            ruc_input = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.ID, "mat-input-3"))
            )
            print("Input field found")
            ruc_input.clear()
            ruc_input.send_keys(ruc)
        except Exception as e:
            print(f"Error finding input field: {str(e)}")
            raise
        time.sleep(10)
        try:
            button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, "button.boton-buscar"))
            )
            print("Button found")
            button.click()
            print("Button clicked")
        except Exception as e:
            print(f"Error with button interaction: {str(e)}")
            raise
        time.sleep(15)
        try:
            button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, "button.boton-buscar"))
            )
            print("Button found")
            button.click()
            print("Button clicked")
        except Exception as e:
            print(f"Error with button interaction: {str(e)}")
            raise
        time.sleep(5)
        screenshot_path = os.path.join(output_dir, 'capture.png')
        driver.save_screenshot(screenshot_path)

    except Exception as e:
        screenshot_path = os.path.join(output_dir, 'capture.png')
        driver.save_screenshot(screenshot_path)

    create_source_table1(document, 'Procesos Judiciales')
    document.add_picture(screenshot_path, width=Inches(6))
    document.add_paragraph('')

    driver.quit()

def scrape_from_soce_incumplidos(document, ruc):
    driver = webdriver.Chrome(service=Service(GeckoDriverManager().install()), options=options)
    try:
        driver.get(url_soce_incumplidos)
        try:
            ruc_input = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.ID, "ruc"))
            )
            print("Input field found")
            ruc_input.clear()
            ruc_input.send_keys(ruc)
        except Exception as e:
            print(f"Error finding input field: {str(e)}")
            raise
        time.sleep(10)
        try:
            button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.ID, "btnBuscar"))
            )
            print("Button found")
            button.click()
            print("Button clicked")
        except Exception as e:
            print(f"Error with button interaction: {str(e)}")
            raise
        time.sleep(15)
        screenshot_path = os.path.join(output_dir, 'capture.png')
        driver.save_screenshot(screenshot_path)

    except Exception as e:
        screenshot_path = os.path.join(output_dir, 'capture.png')
        driver.save_screenshot(screenshot_path)

    create_source_table1(document, 'Búsqueda de no ser contratista incumplido o adjudicatario fallido con el Estado')
    document.add_picture(screenshot_path, width=Inches(6))
    document.add_paragraph('')

    driver.quit()

def scrape_from_contraloria(document, ruc):
    data = get_ruc_data(ruc)
    razon_social = data[0]['razonSocial']
    driver = webdriver.Chrome(service=Service(GeckoDriverManager().install()), options=options)
    try:
        driver.get(url_contraloria)
        try:
            button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.CLASS_NAME, "btn-close"))
            )
            print("Button found")
            button.click()
            print("Button clicked")
        except Exception as e:
            print(f"Error with button interaction: {str(e)}")
            raise
        try:
            ruc_input = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.ID, "txtBuscar123_in"))
            )
            print("Input field found")
            ruc_input.clear()
            ruc_input.send_keys(razon_social)
        except Exception as e:
            print(f"Error finding input field: {str(e)}")
            raise
        
        # Wait for any loading overlay to disappear
        try:
            WebDriverWait(driver, 10).until_not(
                EC.presence_of_element_located((By.CLASS_NAME, "blockUI"))
            )
        except TimeoutException:
            print("Loading overlay did not disappear")
            
        try:
            button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.ID, "btnBuscar_in"))
            )
            # Add a small delay to ensure the button is truly clickable
            time.sleep(2)
            driver.execute_script("arguments[0].click();", button)
            print("Button clicked using JavaScript")
        except Exception as e:
            print(f"Error with button interaction: {str(e)}")
            raise
            
        time.sleep(15)
        screenshot_path = os.path.join(output_dir, 'capture.png')
        driver.save_screenshot(screenshot_path)

    except Exception as e:
        screenshot_path = os.path.join(output_dir, 'capture.png')
        driver.save_screenshot(screenshot_path)

    create_source_table1(document, 'Informes Aprobados')
    document.add_picture(screenshot_path, width=Inches(6))
    document.add_paragraph('')

    driver.quit()

def scrape_from_senescyt(document, ruc):
    driver = webdriver.Chrome(service=Service(GeckoDriverManager().install()), options=options)
    try:
        driver.get(url_senescyt)
        try:
            ruc_input = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.ID, "formPrincipal:identificacion"))
            )
            print("Input field found")
            ruc_input.clear()
            ruc_input.send_keys(ruc)
        except Exception as e:
            print(f"Error finding input field: {str(e)}")
            raise
        time.sleep(15)
        try:
            button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.ID, "formPrincipal:boton-buscar"))
            )
            # Add a small delay to ensure the button is truly clickable
            time.sleep(2)
            driver.execute_script("arguments[0].click();", button)
            print("Button clicked using JavaScript")
        except Exception as e:
            print(f"Error with button interaction: {str(e)}")
            raise

        time.sleep(15)
        screenshot_path = os.path.join(output_dir, 'capture.png')
        driver.save_screenshot(screenshot_path)

    except Exception as e:
        screenshot_path = os.path.join(output_dir, 'capture.png')
        driver.save_screenshot(screenshot_path)

    create_source_table1(document, 'Consulta de formación académica')
    document.add_picture(screenshot_path, width=Inches(6))
    document.add_paragraph('')

    driver.quit()

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'
@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate_report():
    ruc = request.form.get('ruc')

    if not ruc:
        flash('Por favor ingrese un RUC', 'error')
        return redirect(url_for('index'))

    try:
        if validate_ruc(ruc) == 'true':
            document = Document()
            doc_name = f'evidence_report_{ruc}.docx'
            doc_path = os.path.join(output_dir, doc_name)
            add_format(document, ruc)
            create_source_table(document, 'Servicio de Rentas Internas')
            scrape_from_sri(document, ruc)
            create_source_table(document, 'Aduana del Ecuador')
            scrape_from_aduana(document, ruc)
            create_source_table(document, 'Fiscalía')
            scrape_from_fiscalia(document, ruc)
            create_source_table(document, 'Consejo de la Judicatura')
            scrape_from_consejo_judicatura(document, ruc)
            create_source_table(document, 'Servicio Nacional de Contratación Pública (SERCOP)')
            scrape_from_soce_incumplidos(document, ruc)
            create_source_table(document, 'Contraloría General del Estado')
            scrape_from_contraloria(document, ruc)
            identification = add_format2(document, ruc)
            ruc = identification + '001'
            create_source_table(document, 'Servicio de Rentas Internas')
            scrape_from_sri(document, ruc)
            create_source_table(document, 'Fiscalía')
            scrape_from_fiscalia(document, ruc)
            scrape_from_fiscalia(document, identification)
            create_source_table(document, 'Consejo de la Judicatura')
            scrape_from_consejo_judicatura(document, ruc)
            scrape_from_consejo_judicatura(document, identification)
            create_source_table(document, 'Servicio Nacional de Contratación Pública (SERCOP)')
            scrape_from_soce_incumplidos(document, ruc)
            create_source_table(document, 'Ministerio del Interior')
            create_source_table(document, 'Secretaria de Educación superior, ciencia, tecnología e innovación (SENESCYT)')
            scrape_from_senescyt(document, identification)
            document.save(doc_path)

            return send_file(
                doc_path,
                as_attachment=True,
                download_name=doc_name,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            flash('RUC no válido', 'error')
            return redirect(url_for('index'))
    except Exception as e:
        flash(f'Error al generar el reporte: {str(e)}', 'error')
        return redirect(url_for('index'))

if __name__ == '__main__':
    os.makedirs(output_dir, exist_ok=True)
    app.run(host='0.0.0.0', port=5000, debug=True)
